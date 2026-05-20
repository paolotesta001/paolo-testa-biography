from docx import Document
from docx.shared import Pt, Cm, RGBColor

from docx.oxml.ns import qn
from docx.oxml import OxmlElement

doc = Document()

for section in doc.sections:
    section.top_margin = Cm(1.5)
    section.bottom_margin = Cm(1.5)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

style = doc.styles['Normal']
font = style.font
font.name = 'Calibri'
font.size = Pt(10)
font.color.rgb = RGBColor(0x2d, 0x2d, 0x2d)
style.paragraph_format.space_before = Pt(0)
style.paragraph_format.space_after = Pt(0)

DARK = RGBColor(0x1a, 0x1a, 0x1a)
BLUE = RGBColor(0x1a, 0x47, 0x6f)
GRAY = RGBColor(0x55, 0x55, 0x55)
BODY = RGBColor(0x2d, 0x2d, 0x2d)


def right_tab(p):
    tabs = OxmlElement('w:tabs')
    tab = OxmlElement('w:tab')
    tab.set(qn('w:val'), 'right')
    tab.set(qn('w:pos'), '9360')
    tabs.append(tab)
    p._p.get_or_add_pPr().append(tabs)


def add_section_heading(text):
    p = doc.add_paragraph()
    p.space_before = Pt(10)
    p.space_after = Pt(3)
    run = p.add_run(text.upper())
    run.bold = True
    run.font.size = Pt(11)
    run.font.color.rgb = BLUE
    run.font.name = 'Calibri'
    run.italic = True
    pPr = p._p.get_or_add_pPr()
    pBdr = pPr.makeelement(qn('w:pBdr'), {})
    bottom = pBdr.makeelement(qn('w:bottom'), {
        qn('w:val'): 'single', qn('w:sz'): '4',
        qn('w:space'): '1', qn('w:color'): '1a476f'
    })
    pBdr.append(bottom)
    pPr.append(pBdr)


def add_exp_header(company, location, role, date):
    p = doc.add_paragraph()
    p.space_before = Pt(4)
    p.space_after = Pt(0)
    run = p.add_run(company)
    run.bold = True
    run.font.size = Pt(10.5)
    run.font.name = 'Calibri'
    run.font.color.rgb = DARK
    run = p.add_run(' \u2022 ' + location)
    run.font.size = Pt(10)
    run.font.name = 'Calibri'
    run.font.color.rgb = GRAY

    p2 = doc.add_paragraph()
    p2.space_before = Pt(0)
    p2.space_after = Pt(2)
    right_tab(p2)
    run = p2.add_run(role)
    run.font.size = Pt(10)
    run.font.name = 'Calibri'
    run.font.color.rgb = DARK
    run.italic = True
    run = p2.add_run('\t' + date)
    run.font.size = Pt(10)
    run.font.name = 'Calibri'
    run.font.color.rgb = GRAY
    run.italic = True


def add_bullet(text):
    p = doc.add_paragraph(style='List Bullet')
    p.clear()
    p.space_before = Pt(1)
    p.space_after = Pt(1)
    run = p.add_run(text)
    run.font.size = Pt(9.5)
    run.font.name = 'Calibri'
    run.font.color.rgb = BODY


def add_edu_entry(degree, school, date):
    p = doc.add_paragraph()
    p.space_before = Pt(3)
    p.space_after = Pt(1)
    right_tab(p)
    run = p.add_run(degree)
    run.bold = True
    run.font.size = Pt(10)
    run.font.name = 'Calibri'
    run.font.color.rgb = DARK
    run = p.add_run('\t' + date)
    run.font.size = Pt(10)
    run.font.name = 'Calibri'
    run.font.color.rgb = GRAY

    p2 = doc.add_paragraph()
    p2.space_before = Pt(0)
    p2.space_after = Pt(1)
    run = p2.add_run(school)
    run.font.size = Pt(9.5)
    run.font.name = 'Calibri'
    run.font.color.rgb = GRAY


# ===== HEADER =====
name = doc.add_paragraph()
name.space_after = Pt(2)
run = name.add_run('PAOLO TESTA')
run.bold = True
run.font.size = Pt(20)
run.font.name = 'Calibri'
run.font.color.rgb = DARK

contact = doc.add_paragraph()
contact.space_after = Pt(6)
run = contact.add_run(
    'Turin (Italy) \u2022 +393396721513 \u2022 testapaolo2001@gmail.com \u2022 '
    'linkedin.com/in/paolo-testa-4759122b2/ \u2022 github.com/paolotesta001'
)
run.font.size = Pt(9)
run.font.name = 'Calibri'
run.font.color.rgb = GRAY

# ===== SUMMARY (2 lines, systems-engineering focused) =====
add_section_heading('Summary')

p = doc.add_paragraph()
p.space_before = Pt(3)
p.space_after = Pt(2)
run = p.add_run(
    'AI Engineer shipping production-grade multi-agent systems and fault-tolerant LLM platforms. '
    'Hands-on with orchestration, smart model routing, semantic memory (ChromaDB), and end-to-end '
    'MLOps — from data pipelines to Kubernetes deployment with full observability. Background in '
    'cloud infrastructure, robotics, and applied ML across FastAPI, React, Supabase, Docker, Kubernetes, '
    'AWS, and CI/CD.'
)
run.font.size = Pt(10)
run.font.name = 'Calibri'
run.font.color.rgb = BODY

# ===== WORK EXPERIENCE =====
add_section_heading('Work Experience')

# Ex Venture - 5 thematic bullets, each blending platform + Hermes work
add_exp_header('Ex Venture', 'Spain', 'AI Engineer', 'February 2026 \u2013 Present')
for b in [
    'Cost & reliability at scale \u2014 Replaced a \u20ac5,000/month commercial agent platform with an in-house FastAPI orchestrator at '
    '\u20ac1\u20135/month (~99.96% reduction), holding error rate <0.6% across 7 LLM providers under 50 concurrent users via credential '
    'pooling, circuit breakers, and automatic failover.',
    'Smart routing \u2014 task-type and model selection \u2014 Layered a 4-mode task router (CODE / BROWSER / RESEARCH / ORCHESTRATE) '
    'over a per-request model selector choosing between fast-chat, reasoning, and vision models \u2014 91% routing accuracy on a '
    '150-query eval set, with ~40% per-request cost reduction vs. single-model baselines.',
    'Production debugging & MLOps \u2014 Diagnosed a multi-agent deadlock where workers silently hung on provider rate limits '
    '(\u226513 min MTTR against a 600s timeout); replaced indirection with a token-budgeted direct-dispatch path, cutting MTTR to '
    '<30s. Productionized the full stack on Docker, Kubernetes, and GitHub Actions CI/CD with end-to-end observability on '
    'token spend, latency P50/P95, and failover counts.',
    'Memory architecture \u2014 local + semantic \u2014 Implemented a 3-layer memory system: durable user-profile store, 90-day session '
    'log with auto-vacuum and full-text recall, and ChromaDB-backed semantic search \u2014 sub-200ms context loading, '
    'cross-conversation continuity, no external vector-DB dependency for the local tier.',
    'Full-stack shipping beyond the orchestrator \u2014 Extended the OpenClaw platform via external API integrations (OpenAI, '
    'Perplexity, AgentMail), shipped a React/TypeScript/Supabase tracker with real-time sync, and integrated a MobileNetV2 '
    'food-image classifier into a production nutrition platform. Ran a kanban-style 6-worker swarm (researcher, analyst, writer, '
    'reviewer, backend-eng, ops) lifting long-running research throughput ~3\u00d7.',
]:
    add_bullet(b)

# University - AI Engineer Intern
add_exp_header('University', 'Italy', 'AI Engineer Intern', 'September 2024 \u2013 June 2025')
for b in [
    'Engineered an Android application controlling a 20+ DOF humanoid robot with voice-commanded teleoperation and '
    'inverse-kinematics gait control.',
    'Integrated on-device NLP modules to translate voice inputs into real-time robotic maneuvers with high accuracy, '
    'sustaining <50ms command latency over Bluetooth.',
    'Led a 4-person team over 8 months; won 1st place at the National AI Competition (Rome Cup).',
]:
    add_bullet(b)

# Amazon - 1 compact line
add_exp_header('Amazon', 'Turin, Italy', 'Operations / Warehouse Associate', 'May 2021 \u2013 January 2026')
add_bullet('High-volume automated logistics environment; consistently met strict performance and accuracy KPIs.')

# ===== SKILLS (above Projects, reordered for recruiter scanning) =====
add_section_heading('Skills')

skills = [
    ('AI / ML', 'LLM Integration, Multi-Agent Systems, Semantic Memory (ChromaDB), NLP, Computer Vision, PyTorch, TensorFlow, Scikit-Learn, Keras'),
    ('Languages', 'Python, SQL, TypeScript, Rust, Java, Kotlin'),
    ('Cloud & Infra', 'AWS, Azure, OpenStack, Docker, Kubernetes, Terraform, GitHub Actions CI/CD'),
    ('MLOps', 'Model monitoring, observability, token/cost tracking, data pipelines, feature engineering'),
    ('Backend & Data', 'FastAPI, Django, React, Supabase, SQL, NoSQL'),
    ('Tools', 'Git, Linux, n8n, Bolt AI'),
]

for cat, items in skills:
    p = doc.add_paragraph()
    p.space_before = Pt(1)
    p.space_after = Pt(1)
    run = p.add_run(cat + ': ')
    run.bold = True
    run.font.size = Pt(10)
    run.font.name = 'Calibri'
    run.font.color.rgb = DARK
    run = p.add_run(items)
    run.font.size = Pt(9.5)
    run.font.name = 'Calibri'
    run.font.color.rgb = BODY

# ===== PROJECTS (4 substantial, problem-focused descriptions) =====
add_section_heading('Projects')

projects = [
    ('End-to-End AI Application Deployment on AWS',
     'Terraform infrastructure provisioning, Dockerized services, and CI/CD pipelines for scalable backend deployment.',
     'March 2026'),
    ('LLM Workflow Automation',
     'Designed multi-agent LLM workflows with API integrations, robust error handling, and automated task orchestration.',
     'June 2026'),
    ('Time-Series Forecasting (Kaggle)',
     'Forecasting models built on 20 years of time-series data with feature engineering, validation, and hyperparameter tuning.',
     'December 2025'),
    ('OpenStack Load-Balanced Databases',
     'Implemented Octavia load balancer with MySQL replication for fault-tolerant query distribution.',
     '2025'),
]

for proj_name, proj_desc, proj_date in projects:
    p = doc.add_paragraph()
    p.space_before = Pt(3)
    p.space_after = Pt(0)
    right_tab(p)
    run = p.add_run(proj_name)
    run.bold = True
    run.font.size = Pt(10)
    run.font.name = 'Calibri'
    run.font.color.rgb = DARK
    run = p.add_run('\t' + proj_date)
    run.font.size = Pt(9.5)
    run.font.name = 'Calibri'
    run.font.color.rgb = GRAY
    run.italic = True

    p2 = doc.add_paragraph()
    p2.space_before = Pt(0)
    p2.space_after = Pt(1)
    run = p2.add_run(proj_desc)
    run.font.size = Pt(9.5)
    run.font.name = 'Calibri'
    run.font.color.rgb = GRAY

# ===== EDUCATION (cleaner format) =====
add_section_heading('Education')

add_edu_entry(
    'BSc Artificial Intelligence & Computer Science',
    'University of Eastern Piedmont "Amedeo Avogadro", Italy',
    '2024 \u2013 2026'
)
add_edu_entry(
    'Exchange Program \u2014 Artificial Intelligence & Computer Science',
    'NTNU \u2014 Norwegian University of Science and Technology, Norway',
    '2025'
)

# ===== CERTIFICATIONS (with Correlation One moved here) =====
add_section_heading('Certifications')

for cert_name, cert_date in [
    ('AWS Cloud Support Specialist (Honors)', '2025'),
    ('Microsoft Azure AI Solutions (Azure AI Foundry)', '2025'),
    ('Amazon Cloud Support Specialist \u2014 Correlation One', '2025'),
]:
    p = doc.add_paragraph()
    p.space_before = Pt(2)
    p.space_after = Pt(1)
    right_tab(p)
    run = p.add_run(cert_name)
    run.bold = True
    run.font.size = Pt(10)
    run.font.name = 'Calibri'
    run.font.color.rgb = DARK
    run = p.add_run('\t' + cert_date)
    run.font.size = Pt(9.5)
    run.font.name = 'Calibri'
    run.font.color.rgb = GRAY
    run.italic = True

doc.save('Paolo_Testa_CV_v3.docx')
print('Done! Optimized CV saved as Paolo_Testa_CV_v3.docx')
